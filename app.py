from flask import Flask, render_template, request, send_file, jsonify
import pyreadstat
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import io
from datetime import datetime
import os
from config import config

app = Flask(__name__)

# Load configuration based on environment
env = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[env])

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def add_thick_border(paragraph):
    """Add a thick black border below a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '32')  # 24 = 3pt (thick line)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_font(run, size=12, font_name='Times New Roman', bold=False, italic=False):
    """Set font properties for a run"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic

def calculate_distribution_percentages(series):
    """Calculate distribution percentages that sum to 100%"""
    # Drop NaN values before counting
    series_clean = series.dropna()
    value_counts = series_clean.value_counts()
    total = len(series_clean)
    
    # Handle edge case of no data
    if total == 0:
        return {}
    
    # Calculate raw percentages
    raw_percentages = {val: (count / total) * 100 for val, count in value_counts.items()}
    
    # Round to nearest integer
    rounded = {val: round(pct) for val, pct in raw_percentages.items()}
    
    # Adjust to ensure sum is 100
    current_sum = sum(rounded.values())
    diff = 100 - current_sum
    
    if diff != 0 and len(rounded) > 0:
        # Sort by decimal part to determine which values to adjust
        decimals = [(val, raw_percentages[val] - rounded[val]) for val in rounded.keys()]
        decimals.sort(key=lambda x: x[1], reverse=(diff > 0))
        
        # Adjust the values with largest decimal parts
        # Make sure we don't try to adjust more items than we have
        num_adjustments = min(abs(diff), len(decimals))
        for i in range(num_adjustments):
            val = decimals[i][0]
            rounded[val] += 1 if diff > 0 else -1
    
    return rounded

def load_spss_data(file_path):
    """Load SPSS file and return data, metadata"""
    df, meta = pyreadstat.read_sav(file_path)
    return df, meta

def get_question_info(meta):
    """Extract question labels and value labels from metadata"""
    questions = {}
    for col in meta.column_names:
        questions[col] = {
            'label': meta.column_names_to_labels.get(col, col),
            'values': meta.variable_value_labels.get(col, {})
        }
    return questions

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def set_row_height(row, height):
    """Set minimum row height"""
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def set_table_alignment(table, alignment):
    """Set table alignment (center, left, right)"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing jc element if present
    jc_list = tblPr.findall(qn('w:jc'))
    for jc in jc_list:
        tblPr.remove(jc)
    
    # Add new jc element
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)

def parse_apportionment_questions(selected_questions, questions_info):
    """Parse apportionment questions to extract question text and defendant names"""
    if not selected_questions:
        return None, []
    
    # First variable should contain the question and first defendant name
    first_var = selected_questions[0]
    first_label = questions_info[first_var]['label']
    
    # Check if there's a "?" to split on
    if '?' in first_label:
        question_text = first_label.split('?')[0] + '?'
        first_defendant = first_label.split('?')[1].strip()
    else:
        # If no "?", use the whole thing as question and first variable as defendant
        question_text = "What percentage of fault, if any, do you assign to the following?"
        first_defendant = first_label.strip()
    
    # Build list of defendants
    defendants = [{'name': first_defendant, 'variable': first_var}]
    
    # Remaining variables are just defendant names
    for var in selected_questions[1:]:
        defendant_name = questions_info[var]['label'].strip()
        defendants.append({'name': defendant_name, 'variable': var})
    
    return question_text, defendants

def parse_damages_questions(selected_questions, questions_info):
    """Parse damages questions to extract question text and damage categories"""
    if not selected_questions:
        return None, []
    
    # First variable should contain the question and first damage category
    first_var = selected_questions[0]
    first_label = questions_info[first_var]['label']
    
    # Check if there's a "?" to split on
    if '?' in first_label:
        question_text = first_label.split('?')[0] + '?'
        first_category = first_label.split('?')[1].strip()
    else:
        # If no "?", use the whole thing as question and first variable as category
        question_text = "What amount of damages, if any, would you award for the following?"
        first_category = first_label.strip()
    
    # Build list of damage categories
    categories = [{'name': first_category, 'variable': first_var}]
    
    # Remaining variables are just category names
    for var in selected_questions[1:]:
        category_name = questions_info[var]['label'].strip()
        categories.append({'name': category_name, 'variable': var})
    
    return question_text, categories

def create_apportionment_grid(doc, df, defendant_name, defendant_var, letter):
    """Create apportionment grid for a single defendant"""
    series_clean = df[defendant_var].dropna()
    
    if len(series_clean) == 0:
        return
    
    # Calculate statistics
    mean_val = series_clean.mean()
    median_val = series_clean.median()
    min_val = series_clean.min()
    max_val = series_clean.max()
    
    # Get value counts sorted by value (descending)
    value_counts = series_clean.value_counts().sort_index(ascending=False)
    total_n = len(series_clean)
    
    # Create table: 2 columns, rows = header + mean + median + range + amount header + value rows
    num_value_rows = len(value_counts)
    num_rows = 5 + num_value_rows  # header + mean + median + range + amount + values
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths and row heights forcefully
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(1)
        set_row_height(row, 360)  # Minimum row height
    
    # Center the table
    set_table_alignment(table, 'center')
    
    # Row 0: Defendant Name (header)
    header_cell = table.rows[0].cells[0]
    header_cell.merge(table.rows[0].cells[1])
    header_cell.text = defendant_name
    for paragraph in header_cell.paragraphs:
        for run in paragraph.runs:
            set_font(run, size=12, bold=True)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(header_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 1: Mean
    mean_cell = table.rows[1].cells[0]
    mean_cell.merge(table.rows[1].cells[1])
    mean_para = mean_cell.paragraphs[0]
    mean_run = mean_para.add_run('Mean: ')
    set_font(mean_run, size=12, bold=True)
    val_run = mean_para.add_run(f'{mean_val:.0f}%')
    set_font(val_run, size=12, bold=True)
    mean_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(mean_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 2: Median
    median_cell = table.rows[2].cells[0]
    median_cell.merge(table.rows[2].cells[1])
    median_para = median_cell.paragraphs[0]
    median_run = median_para.add_run('Median: ')
    set_font(median_run, size=12, bold=True)
    val_run = median_para.add_run(f'{median_val:.0f}%')
    set_font(val_run, size=12, bold=True)
    median_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(median_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 3: Range
    range_cell = table.rows[3].cells[0]
    range_cell.merge(table.rows[3].cells[1])
    range_para = range_cell.paragraphs[0]
    range_run = range_para.add_run('Range: ')
    set_font(range_run, size=12, bold=True)
    val_run = range_para.add_run(f'{min_val:.0f}% - {max_val:.0f}%')
    set_font(val_run, size=12, bold=True)
    range_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(range_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 4: Amount header
    amount_left = table.rows[4].cells[0]
    amount_right = table.rows[4].cells[1]
    
    amount_left_para = amount_left.paragraphs[0]
    amount_run = amount_left_para.add_run('Amount:')
    set_font(amount_run, size=12, bold=True)
    amount_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(amount_left, top=True, left=True, bottom=True, right=True)
    
    amount_right_para = amount_right.paragraphs[0]
    amount_right_para.add_run('(')
    n_run = amount_right_para.add_run('n')
    set_font(n_run, size=12, italic=True, bold=True)
    last_run = amount_right_para.add_run(f' = {total_n})')
    set_font(last_run, size=12, bold=True)
    amount_right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    amount_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(amount_right, top=True, left=True, bottom=True, right=True)
    
    # Value rows
    for idx, (value, count) in enumerate(value_counts.items(), start=5):
        left_cell = table.rows[idx].cells[0]
        right_cell = table.rows[idx].cells[1]
        
        # Left: percentage value
        left_cell.text = f'{value:.0f}%'
        for run in left_cell.paragraphs[0].runs:
            set_font(run, size=12)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(left_cell, top=True, left=True, bottom=True, right=True)
        
        # Right: (n = count)
        right_para = right_cell.paragraphs[0]
        right_para.add_run('(')
        n_run = right_para.add_run('n')
        set_font(n_run, size=12, italic=True)
        last_run = right_para.add_run(f' = {count})')
        set_font(last_run, size=12)
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(right_cell, top=True, left=True, bottom=True, right=True)
    
    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

def create_damages_grid(doc, df, category_name, category_var, letter):
    """Create damages grid for a single category"""
    series_clean = df[category_var].dropna()
    
    if len(series_clean) == 0:
        return
    
    # Calculate statistics
    mean_val = series_clean.mean()
    median_val = series_clean.median()
    min_val = series_clean.min()
    max_val = series_clean.max()
    
    # Get value counts sorted by value (descending)
    value_counts = series_clean.value_counts().sort_index(ascending=False)
    total_n = len(series_clean)
    
    # Create table: 2 columns, rows = header + mean + median + range + amount header + value rows
    num_value_rows = len(value_counts)
    num_rows = 5 + num_value_rows  # header + mean + median + range + amount + values
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths and row heights forcefully
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(1)
        set_row_height(row, 360)  # Minimum row height
    
    # Center the table
    set_table_alignment(table, 'center')
    
    # Row 0: Category Name (header)
    header_cell = table.rows[0].cells[0]
    header_cell.merge(table.rows[0].cells[1])
    header_cell.text = category_name
    for paragraph in header_cell.paragraphs:
        for run in paragraph.runs:
            set_font(run, size=12, bold=True)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(header_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 1: Mean
    mean_cell = table.rows[1].cells[0]
    mean_cell.merge(table.rows[1].cells[1])
    mean_para = mean_cell.paragraphs[0]
    mean_run = mean_para.add_run('Mean: ')
    set_font(mean_run, size=12, bold=True)
    val_run = mean_para.add_run(f'${mean_val:,.0f}')
    set_font(val_run, size=12, bold=True)
    mean_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(mean_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 2: Median
    median_cell = table.rows[2].cells[0]
    median_cell.merge(table.rows[2].cells[1])
    median_para = median_cell.paragraphs[0]
    median_run = median_para.add_run('Median: ')
    set_font(median_run, size=12, bold=True)
    val_run = median_para.add_run(f'${median_val:,.0f}')
    set_font(val_run, size=12, bold=True)
    median_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(median_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 3: Range
    range_cell = table.rows[3].cells[0]
    range_cell.merge(table.rows[3].cells[1])
    range_para = range_cell.paragraphs[0]
    range_run = range_para.add_run('Range: ')
    set_font(range_run, size=12, bold=True)
    val_run = range_para.add_run(f'${min_val:,.0f} - ${max_val:,.0f}')
    set_font(val_run, size=12, bold=True)
    range_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(range_cell, top=True, left=True, bottom=True, right=True)
    
    # Row 4: Amount header
    amount_left = table.rows[4].cells[0]
    amount_right = table.rows[4].cells[1]
    
    amount_left_para = amount_left.paragraphs[0]
    amount_run = amount_left_para.add_run('Amount:')
    set_font(amount_run, size=12, bold=True)
    amount_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(amount_left, top=True, left=True, bottom=True, right=True)
    
    amount_right_para = amount_right.paragraphs[0]
    amount_right_para.add_run('(')
    n_run = amount_right_para.add_run('n')
    set_font(n_run, size=12, italic=True, bold=True)
    last_run = amount_right_para.add_run(f' = {total_n})')
    set_font(last_run, size=12, bold=True)
    amount_right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    amount_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(amount_right, top=True, left=True, bottom=True, right=True)
    
    # Value rows
    for idx, (value, count) in enumerate(value_counts.items(), start=5):
        left_cell = table.rows[idx].cells[0]
        right_cell = table.rows[idx].cells[1]
        
        # Left: dollar value
        left_cell.text = f'${value:,.0f}'
        for run in left_cell.paragraphs[0].runs:
            set_font(run, size=12)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(left_cell, top=True, left=True, bottom=True, right=True)
        
        # Right: (n = count)
        right_para = right_cell.paragraphs[0]
        right_para.add_run('(')
        n_run = right_para.add_run('n')
        set_font(n_run, size=12, italic=True)
        last_run = right_para.add_run(f' = {count})')
        set_font(last_run, size=12)
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(right_cell, top=True, left=True, bottom=True, right=True)
    
    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

def create_apportionment_section(doc, df, meta, selected_questions, question_number):
    """Create the apportionment section with main question and defendant grids"""
    if not selected_questions:
        return question_number
    
    questions_info = get_question_info(meta)
    
    # Parse questions to get question text and defendants
    question_text, defendants = parse_apportionment_questions(selected_questions, questions_info)
    
    if not question_text or not defendants:
        return question_number
    
    # Add "Apportionment" header (bold and underlined)
    apportionment_header = doc.add_paragraph()
    run = apportionment_header.add_run('Apportionment')
    set_font(run, size=12, bold=True)
    run.underline = True
    apportionment_header.paragraph_format.space_after = Pt(12)
    
    # Add main question with number (BOLD)
    main_question = doc.add_paragraph()
    run = main_question.add_run(f'{question_number}. {question_text}')
    set_font(run, size=12, bold=True)
    main_question.paragraph_format.space_after = Pt(12)
    
    # Create grid for each defendant
    letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
    
    for idx, defendant in enumerate(defendants):
        if idx < len(letters):
            letter = letters[idx]
            # Add defendant header with letter (NOT BOLD, with tab indent)
            defendant_header = doc.add_paragraph()
            defendant_header.paragraph_format.left_indent = Inches(0.5)
            run = defendant_header.add_run(f'{letter}. {defendant["name"]}')
            set_font(run, size=12, bold=False)
            defendant_header.paragraph_format.space_after = Pt(6)
            
            # Create the grid for this defendant
            create_apportionment_grid(doc, df, defendant['name'], defendant['variable'], letter)
    
    return question_number + 1

def create_damages_section(doc, df, meta, selected_questions, question_number):
    """Create the damages section with main question and category grids"""
    if not selected_questions:
        return question_number
    
    questions_info = get_question_info(meta)
    
    # Parse questions to get question text and categories
    question_text, categories = parse_damages_questions(selected_questions, questions_info)
    
    if not question_text or not categories:
        return question_number
    
    # Add "Damages" header (bold and underlined)
    damages_header = doc.add_paragraph()
    run = damages_header.add_run('Damages')
    set_font(run, size=12, bold=True)
    run.underline = True
    damages_header.paragraph_format.space_after = Pt(12)
    
    # Add main question with number (BOLD)
    main_question = doc.add_paragraph()
    run = main_question.add_run(f'{question_number}. {question_text}')
    set_font(run, size=12, bold=True)
    main_question.paragraph_format.space_after = Pt(12)
    
    # Create grid for each category
    letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
    
    for idx, category in enumerate(categories):
        if idx < len(letters):
            letter = letters[idx]
            # Add category header with letter (NOT BOLD, with tab indent)
            category_header = doc.add_paragraph()
            category_header.paragraph_format.left_indent = Inches(0.5)
            run = category_header.add_run(f'{letter}. {category["name"]}')
            set_font(run, size=12, bold=False)
            category_header.paragraph_format.space_after = Pt(6)
            
            # Create the grid for this category
            create_damages_grid(doc, df, category['name'], category['variable'], letter)
    
    return question_number + 1

def create_argument_grid(doc, df, meta, selected_questions, party_name, party_type):
    """Create argument grid table for plaintiff or defendant"""
    questions_info = get_question_info(meta)
    ordered_questions = [q for q in meta.column_names if q in selected_questions]
    
    if not ordered_questions:
        return
    
    # Add question header
    question_header = doc.add_paragraph()
    question_text = f'How persuasive were each of the following arguments made on behalf of the {party_type}, {party_name}?'
    run = question_header.add_run(question_text)
    set_font(run, size=12, bold=True)
    question_header.paragraph_format.space_after = Pt(6)
    
    # Create table: 6 columns, rows = 1 header + number of selected questions
    num_rows = 1 + len(ordered_questions)
    table = doc.add_table(rows=num_rows, cols=6)
    table.style = 'Table Grid'
    
    # Set column widths for each cell
    widths = [Inches(0.5), Inches(3.0), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Mean', f'{party_type} Arguments', 'Disagree strongly', 'Disagree somewhat', 'Agree somewhat', 'Agree strongly']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        font_size = 12 if i < 2 else 10
        # Bold and center the header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=font_size, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add padding to paragraph
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
        # Add borders
        set_cell_border(cell, top=True, left=True, bottom=True, right=True)
    
    # Data rows
    for idx, question_var in enumerate(ordered_questions, start=1):
        if question_var not in df.columns:
            continue
        
        row = table.rows[idx]
        cells = row.cells
        
        # Calculate mean (numeric values 1-4)
        series_clean = df[question_var].dropna()
        if len(series_clean) > 0:
            mean_val = series_clean.mean()
        else:
            mean_val = 0
        
        # Calculate percentages for each response option
        percentages = calculate_distribution_percentages(df[question_var])
        
        # Get the argument text (question label)
        argument_text = questions_info[question_var]['label']
        
        # Fill in cells
        # Column 0: Mean (bold)
        cells[0].text = f'{mean_val:.2f}'
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=12, bold=True)
        # Add padding
        cells[0].paragraphs[0].paragraph_format.space_before = Pt(6)
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(6)
        set_cell_border(cells[0], top=True, left=True, bottom=True, right=True)
        
        # Column 1: Argument text
        cells[1].text = argument_text
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=12)
        # Add padding
        cells[1].paragraphs[0].paragraph_format.space_before = Pt(6)
        cells[1].paragraphs[0].paragraph_format.space_after = Pt(6)
        set_cell_border(cells[1], top=True, left=True, bottom=True, right=True)
        
        # Columns 2-5: Response percentages (values 1-4) - bold
        for i, value in enumerate([1, 2, 3, 4], start=2):
            pct = percentages.get(value, 0)
            cells[i].text = f'{pct}%'
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[i].paragraphs[0].runs:
                set_font(run, size=12, bold=True)
            # Add padding
            cells[i].paragraphs[0].paragraph_format.space_before = Pt(6)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(6)
            set_cell_border(cells[i], top=True, left=True, bottom=True, right=True)
    
    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

def create_combined_document(df, meta, mc_questions, plaintiff_questions, defendant_questions, 
                            apportionment_questions, damages_questions, questionnaire_num, case_id, 
                            plaintiff_name, defendant_name, custom_date):
    """Create document with multiple choice questions, argument grids, apportionment, and damages"""
    doc = Document()
    
    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header: "Posted Questionnaire #"
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f'Posted Questionnaire #{questionnaire_num}')
    set_font(run, size=14, bold=True)
    header.paragraph_format.space_after = Pt(0)
    
    # "(After Closings and Rebuttal)"
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheader.add_run('(After Closings and Rebuttal)')
    set_font(run, size=12)
    subheader.paragraph_format.space_before = Pt(0)
    subheader.paragraph_format.space_after = Pt(0)
    
    # "(n = X)"
    n_total = len(df)
    n_line = doc.add_paragraph()
    n_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paren_run = n_line.add_run('(')
    set_font(paren_run, size=12)
    n_run = n_line.add_run('n')
    set_font(n_run, size=12, italic=True)
    equals_run = n_line.add_run(f' = {n_total})')
    set_font(equals_run, size=12)
    n_line.paragraph_format.space_before = Pt(0)
    n_line.paragraph_format.space_after = Pt(6)
    
    # Thick black line separator
    separator = doc.add_paragraph()
    add_thick_border(separator)
    separator.paragraph_format.space_after = Pt(12)
    
    # Get question info
    questions_info = get_question_info(meta)
    
    # Track question numbering
    current_question_num = 1
    
    # Multiple Choice Section
    if mc_questions:
        ordered_questions = [q for q in meta.column_names if q in mc_questions]
        
        for idx, question_var in enumerate(ordered_questions):
            if question_var not in df.columns:
                continue
                
            question_info = questions_info[question_var]
            question_label = question_info['label']
            value_labels = question_info['values']
            
            # Add question number and text (bolded)
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f'{current_question_num}.  {question_label}')
            set_font(question_run, size=12, bold=True)
            question_para.paragraph_format.space_after = Pt(12)
            
            # Calculate distribution with proper rounding
            percentages = calculate_distribution_percentages(df[question_var])
            value_counts = df[question_var].value_counts()
            
            # Determine if we should display inline (2 or fewer options) or stacked (3+ options)
            num_options = len(value_labels)
            inline_display = num_options <= 2
            
            if inline_display:
                # Display all responses on one line with triple spacing between them (with tab indent)
                response_para = doc.add_paragraph()
                response_para.paragraph_format.left_indent = Inches(0.5)
                response_text_parts = []
                
                for value_code in sorted(value_labels.keys()):
                    if value_code in value_counts.index:
                        count = value_counts[value_code]
                        pct = percentages.get(value_code, 0)
                        response_label = value_labels[value_code]
                        response_text_parts.append((pct, response_label, count))
                
                # Add all parts with proper spacing
                for i, (pct, response_label, count) in enumerate(response_text_parts):
                    if i > 0:
                        response_para.add_run('                 ')  # triple tab spacing
                    
                    # Add text without quotation marks
                    text_run = response_para.add_run(f'{pct}%  {response_label}  (')
                    set_font(text_run, size=12)
                    italic_run = response_para.add_run('n')
                    set_font(italic_run, size=12, italic=True)
                    end_run = response_para.add_run(f' = {count})')
                    set_font(end_run, size=12)
                
                response_para.paragraph_format.space_after = Pt(24)
            else:
                # Display each response on its own line (with tab indent)
                for value_code in sorted(value_labels.keys()):
                    if value_code in value_counts.index:
                        count = value_counts[value_code]
                        pct = percentages.get(value_code, 0)
                        response_label = value_labels[value_code]
                        
                        response_para = doc.add_paragraph()
                        response_para.paragraph_format.left_indent = Inches(0.5)
                        
                        # Add space before single-digit percentages for alignment
                        pct_text = f'  {pct}%' if pct < 10 else f'{pct}%'
                        
                        text_run = response_para.add_run(f'{pct_text}  {response_label}  (')
                        set_font(text_run, size=12)
                        italic_run = response_para.add_run('n')
                        set_font(italic_run, size=12, italic=True)
                        end_run = response_para.add_run(f' = {count})')
                        set_font(end_run, size=12)
                        response_para.paragraph_format.space_after = Pt(3)
                
                # Add spacing after the last response option
                if len(value_labels) > 0:
                    response_para.paragraph_format.space_after = Pt(24)
            
            current_question_num += 1
    
    # Apportionment Section
    if apportionment_questions:
        current_question_num = create_apportionment_section(doc, df, meta, apportionment_questions, current_question_num)
    
    # Damages Section
    if damages_questions:
        current_question_num = create_damages_section(doc, df, meta, damages_questions, current_question_num)
    
    # Plaintiff Argument Grid
    if plaintiff_questions and plaintiff_name:
        create_argument_grid(doc, df, meta, plaintiff_questions, plaintiff_name, 'Plaintiff')
    
    # Defendant Argument Grid
    if defendant_questions and defendant_name:
        create_argument_grid(doc, df, meta, defendant_questions, defendant_name, 'Defendant')
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    
    # Create a table in the footer for left, center, and right alignment
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    footer_table.autofit = False
    
    # Set column widths
    footer_table.rows[0].cells[0].width = Inches(2.17)
    footer_table.rows[0].cells[1].width = Inches(2.16)
    footer_table.rows[0].cells[2].width = Inches(2.17)
    
    # Left cell - use custom_date instead of datetime.now()
    left_cell = footer_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_run1 = left_para.add_run(f'Q{questionnaire_num} ({custom_date})\n')
    set_font(left_run1, size=8)
    left_run2 = left_para.add_run(case_id)
    set_font(left_run2, size=8)
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Center cell (page number)
    center_cell = footer_table.rows[0].cells[1]
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run = center_para.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, size=8)
    
    # Right cell
    right_cell = footer_table.rows[0].cells[2]
    right_para = right_cell.paragraphs[0]
    right_run1 = right_para.add_run('PRIVILEGED & CONFIDENTIAL\n')
    set_font(right_run1, size=8)
    right_run2 = right_para.add_run('ATTORNEY WORK PRODUCT')
    set_font(right_run2, size=8)
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Remove table borders
    for row in footer_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    return doc

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and return available questions"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.endswith('.sav'):
        return jsonify({'error': 'Please upload an SPSS (.sav) file'}), 400
    
    try:
        # Save file temporarily
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        
        # Load SPSS data
        df, meta = load_spss_data(file_path)
        
        # Get question information
        questions_info = get_question_info(meta)
        
        # Prepare question list for frontend
        questions_list = []
        for var_name in meta.column_names:
            questions_list.append({
                'variable': var_name,
                'label': questions_info[var_name]['label'],
                'num_values': len(questions_info[var_name]['values'])
            })
        
        return jsonify({
            'success': True,
            'filename': file.filename,
            'questions': questions_list,
            'total_responses': len(df)
        })
    
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/generate', methods=['POST'])
def generate_document():
    """Generate the Word document"""
    try:
        data = request.json
        filename = data.get('filename')
        mc_questions = data.get('mc_questions', [])
        plaintiff_questions = data.get('plaintiff_questions', [])
        defendant_questions = data.get('defendant_questions', [])
        apportionment_questions = data.get('apportionment_questions', [])
        damages_questions = data.get('damages_questions', [])
        questionnaire_num = data.get('questionnaire_num', '1')
        case_id = data.get('case_id', '')
        plaintiff_name = data.get('plaintiff_name', '')
        defendant_name = data.get('defendant_name', '')
        custom_date = data.get('custom_date', '')  # Get custom date from request
        
        if not filename:
            return jsonify({'error': 'Missing filename'}), 400
        
        if not mc_questions and not plaintiff_questions and not defendant_questions and not apportionment_questions and not damages_questions:
            return jsonify({'error': 'Please select at least one question'}), 400
        
        # Load the file
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found. Please upload again.'}), 400
        
        df, meta = load_spss_data(file_path)
        
        # Create document with custom date
        doc = create_combined_document(df, meta, mc_questions, plaintiff_questions, defendant_questions,
                                      apportionment_questions, damages_questions, questionnaire_num, 
                                      case_id, plaintiff_name, defendant_name, custom_date)
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Posted_Q{questionnaire_num}_{case_id.replace(" ", "_")}.docx'
        )
    
    except Exception as e:
        return jsonify({'error': f'Error generating document: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)